import os
import copy
import qrcode
import pypdf
from datetime import datetime
from io import BytesIO
from django.core.files.base import ContentFile
from django.conf import settings
from reportlab.lib import colors
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas as pdfcanvas
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# Register fonts
FONT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'fonts')
pdfmetrics.registerFont(TTFont('Poppins-Bold', os.path.join(FONT_DIR, 'Poppins-Bold.ttf')))
pdfmetrics.registerFont(TTFont('Poppins-Regular', os.path.join(FONT_DIR, 'Poppins-Regular.ttf')))
pdfmetrics.registerFont(TTFont('LibreBaskerville-Regular', os.path.join(FONT_DIR, 'LibreBaskerville-Regular.ttf')))
pdfmetrics.registerFont(TTFont('LibreBaskerville-Bold', os.path.join(FONT_DIR, 'LibreBaskerville-Bold.ttf')))

# ─── python-docx imports ──────────────────────────────────────────────────────
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from lxml import etree

# ─── Brand colors ─────────────────────────────────────────────────────────────
BRAND_NAVY   = colors.HexColor('#041321')   # Dark navy (matches Word color)
BRAND_BLUE   = colors.HexColor('#1F3F70')   # Mid blue  (matches Word color)

# ─── Uzbek months ─────────────────────────────────────────────────────────────
UZBEK_MONTHS = {
    1: "yanvar", 2: "fevral", 3: "mart", 4: "aprel", 5: "may", 6: "iyun",
    7: "iyul", 8: "avgust", 9: "sentabr", 10: "oktabr", 11: "noyabr", 12: "dekabr"
}


def _get_template_path(filename):
    """Return absolute path to a template file, checking multiple locations."""
    paths = [
        os.path.join(settings.BASE_DIR, '..', filename),
        os.path.join(r'D:\SES', filename),
        os.path.join(os.path.dirname(settings.BASE_DIR), filename),
    ]
    for p in paths:
        if os.path.exists(p):
            return os.path.abspath(p)
    raise FileNotFoundError(f"Template file '{filename}' not found. Checked: {paths}")


# ─────────────────────────────────────────────────────────────────────────────
# 1. QR Code generator
# ─────────────────────────────────────────────────────────────────────────────
def _generate_qr(certificate):
    """Generate QR code image and save to certificate.qr_code_image."""
    verify_url = f"https://shahar-ses.uz/verify/{certificate.certificate_id}"

    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_M,
        box_size=10,
        border=2,
    )
    qr.add_data(verify_url)
    qr.make(fit=True)

    qr_img = qr.make_image(fill_color="#041321", back_color="white")
    qr_io = BytesIO()
    qr_img.save(qr_io, format='PNG')
    qr_io.seek(0)

    qr_file_name = f"{certificate.certificate_id}_qr.png"
    certificate.qr_code_image.save(
        qr_file_name,
        ContentFile(qr_io.read()),
        save=False
    )
    return verify_url  # return URL so PDF can embed it


# ─────────────────────────────────────────────────────────────────────────────
# 2. PDF Certificate generator (overlay on template PDF)
# ─────────────────────────────────────────────────────────────────────────────
def _generate_pdf(certificate):
    """Merge data overlay onto PDF template and save to certificate.pdf_file."""
    template_path = _get_template_path('Sterilizatsiya sertifikat.pdf')

    # Page size (landscape A4 in points)
    w, h = 842.25, 595.5

    # Student info
    father = getattr(certificate.student, 'father_name', '')
    if father:
        student_name = f"{certificate.student.first_name} {certificate.student.last_name} {father}".strip()
    else:
        student_name = (
            certificate.student.get_full_name()
            or f"{certificate.student.first_name} {certificate.student.last_name}".strip()
            or certificate.student.username
        )

    issued_date = certificate.issued_at if certificate.issued_at else datetime.now()
    expiry_date = certificate.expires_at

    # ── Overlay canvas ────────────────────────────────────────────────────────
    overlay_io = BytesIO()
    c = pdfcanvas.Canvas(overlay_io, pagesize=(w, h))

    # Sertifikat ID
    # "Ushbu" x1=335.8, "sonli" x0=483.3 => markaz=409.6, rl_y=379
    c.setFont("Poppins-Bold", 13)
    c.setFillColor(BRAND_NAVY)
    c.drawCentredString(409.6, 358, certificate.certificate_id)

    # Talaba ismi
    # chiziq x0=155, "ga" x0=689.6 => markaz=422.3, rl_y=335
    c.setFont("Poppins-Bold", 20)
    c.setFillColor(BRAND_BLUE)
    c.drawCentredString(422.3, 316, f"{student_name.upper()}GA")

    # Boshlanish sanasi (pdfplumber aniq koordinatalar)
    # Start open-quote x0=236.0, close-quote x1=276.2 => kun markazi=256
    # "dan" x0=373.1 => oy markazi=(276+373)/2=324.5
    c.setFont("LibreBaskerville-Bold", 13)
    c.setFillColor(BRAND_NAVY)
    c.drawCentredString(257, 195.6, str(issued_date.day))
    c.setFont("LibreBaskerville-Bold", 13)
    c.drawCentredString(328, 195.6, f"{UZBEK_MONTHS[issued_date.month]}dan")

    # Tugash sanasi (pdfplumber aniq koordinatalar)
    # End open-quote x0=524.1, close-quote x1=564.3 => kun markazi=544
    # "gacha" x0=661.1 => oy markazi=(564+661)/2=612.5
    c.setFont("LibreBaskerville-Bold", 13)
    c.drawCentredString(549, 195.6, str(expiry_date.day))
    c.setFont("LibreBaskerville-Bold", 13)
    c.drawCentredString(618, 195.6, f"{UZBEK_MONTHS[expiry_date.month]}gacha")

    # QR kod - chiziq ustiga joylashtirish (o'ngroqqa ancha surilgan: x=575)
    qr_path = certificate.qr_code_image.path if certificate.qr_code_image else None
    if qr_path and os.path.exists(qr_path):
        c.drawImage(
            ImageReader(qr_path),
            x=575, y=90,
            width=72, height=72,
            preserveAspectRatio=True
        )
    else:
        verify_url = f"https://shahar-ses.uz/verify/{certificate.certificate_id}"
        qr = qrcode.QRCode(version=1, error_correction=qrcode.constants.ERROR_CORRECT_M,
                           box_size=10, border=2)
        qr.add_data(verify_url)
        qr.make(fit=True)
        qr_img = qr.make_image(fill_color="#041321", back_color="white")
        qr_fb = BytesIO()
        qr_img.save(qr_fb, format='PNG')
        qr_fb.seek(0)
        c.drawImage(ImageReader(qr_fb), x=575, y=90, width=72, height=72,
                    preserveAspectRatio=True)

    # QR kod ostiga havolasini yozish (o'ngroqqa ancha surilgan: x=611)
    c.setFont("Poppins-Regular", 7)
    c.setFillColor(BRAND_NAVY)
    c.drawCentredString(611, 80, f"shahar-ses.uz/verify/{certificate.certificate_id}")

    c.showPage()
    c.save()
    overlay_io.seek(0)

    # ── Merge overlay with template ───────────────────────────────────────────
    template_pdf = pypdf.PdfReader(template_path)
    template_page = template_pdf.pages[0]

    # Clean the template page by removing the underline behind the QR code on-the-fly
    try:
        contents = template_page["/Contents"]
        if isinstance(contents, pypdf.generic.ArrayObject):
            # It's an array of streams
            for obj in contents:
                resolved_obj = obj.get_object()
                if hasattr(resolved_obj, "get_data") and hasattr(resolved_obj, "set_data"):
                    data = resolved_obj.get_data()
                    import re as std_re
                    pattern = rb'(?:/gs2\s*gs\s*)?/Image\s*<<\s*/MCID\s*19\s*>>\s*BDC\s*492\.653\s*117\.74017\s*148\.51779\s*0\.75\s*re\s*f\s*EMC'
                    new_data, count = std_re.subn(pattern, b'', data)
                    if count > 0:
                        resolved_obj.set_data(new_data)
                    else:
                        new_data, count = std_re.subn(rb'492\.653\s*117\.74017\s*148\.51779\s*0\.75\s*re\s*f', b'', data)
                        if count > 0:
                            resolved_obj.set_data(new_data)
        elif hasattr(contents, "get_data") and hasattr(contents, "set_data"):
            # It's a single stream
            data = contents.get_data()
            import re as std_re
            pattern = rb'(?:/gs2\s*gs\s*)?/Image\s*<<\s*/MCID\s*19\s*>>\s*BDC\s*492\.653\s*117\.74017\s*148\.51779\s*0\.75\s*re\s*f\s*EMC'
            new_data, count = std_re.subn(pattern, b'', data)
            if count > 0:
                contents.set_data(new_data)
            else:
                new_data, count = std_re.subn(rb'492\.653\s*117\.74017\s*148\.51779\s*0\.75\s*re\s*f', b'', data)
                if count > 0:
                    contents.set_data(new_data)
    except Exception as e:
        import logging
        logging.getLogger(__name__).warning(f"Could not remove line from PDF template: {e}")

    overlay_pdf = pypdf.PdfReader(overlay_io)
    overlay_page = overlay_pdf.pages[0]

    template_page.merge_page(overlay_page)

    writer = pypdf.PdfWriter()
    writer.add_page(template_page)

    pdf_output_io = BytesIO()
    writer.write(pdf_output_io)
    pdf_output_io.seek(0)

    pdf_file_name = f"{certificate.certificate_id}.pdf"
    certificate.pdf_file.save(
        pdf_file_name,
        ContentFile(pdf_output_io.read()),
        save=False
    )


# ─────────────────────────────────────────────────────────────────────────────
# 3. Word Certificate generator (fill text boxes in .docx template)
# ─────────────────────────────────────────────────────────────────────────────
def _generate_word(certificate):
    """Fill Word (.docx) template with certificate data and save to media."""

    template_path = _get_template_path('Sterilizatsiya sertifikat word.docx')

    doc = DocxDocument(template_path)

    # Student info
    father = getattr(certificate.student, 'father_name', '')
    if father:
        student_name = f"{certificate.student.first_name} {certificate.student.last_name} {father}".strip()
    else:
        student_name = (
            certificate.student.get_full_name()
            or f"{certificate.student.first_name} {certificate.student.last_name}".strip()
            or certificate.student.username
        )

    issued_date = certificate.issued_at if certificate.issued_at else datetime.now()
    expiry_date = certificate.expires_at

    issued_day   = str(issued_date.day)
    issued_month = UZBEK_MONTHS[issued_date.month]
    issued_year  = str(issued_date.year)
    expiry_day   = str(expiry_date.day)
    expiry_month = UZBEK_MONTHS[expiry_date.month]
    expiry_year  = str(expiry_date.year)

    W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    WPS_NS = 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'

    root = etree.fromstring(doc.element.body.xml.encode())
    txbx_list = root.findall(f'.//{{{WPS_NS}}}txbx')

    def set_txbx_text(txbx, new_text, preserve_first_run=True):
        """Replace all text in a text box with new_text, preserving first run formatting."""
        runs = txbx.findall(f'.//{{{W_NS}}}r')
        if not runs:
            return

        # Get the first run's XML properties for formatting preservation
        first_run = runs[0]
        first_rpr = first_run.find(f'{{{W_NS}}}rPr')

        # Remove all existing runs from their parent paragraphs
        paras = txbx.findall(f'.//{{{W_NS}}}p')
        for para in paras:
            for run in para.findall(f'{{{W_NS}}}r'):
                para.remove(run)

        # Write new text into first paragraph
        if paras:
            para = paras[0]
            new_run = etree.SubElement(para, f'{{{W_NS}}}r')
            if first_rpr is not None:
                new_run.insert(0, copy.deepcopy(first_rpr))
            t_el = etree.SubElement(new_run, f'{{{W_NS}}}t')
            t_el.text = new_text
            if new_text.startswith(' ') or new_text.endswith(' '):
                t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    # ── Sertifikat 1 (chap/birinchi nusxa) ───────────────────────────────────
    # TextBox[9]  → Sertifikat ID
    # TextBox[15] → Talaba ismi
    # TextBox[19] → Boshlanish yili
    # TextBox[22] → Boshlanish kuni
    # TextBox[24] → Boshlanish oyi
    # TextBox[28] → "dan" (o'zgartirmasdan qoldiramiz)
    # TextBox[30] → Tugash yili
    # TextBox[33] → Tugash kuni
    # TextBox[35] → Tugash oyi

    if len(txbx_list) > 9:
        set_txbx_text(txbx_list[9], certificate.certificate_id)
    if len(txbx_list) > 15:
        set_txbx_text(txbx_list[15], student_name.upper())
    if len(txbx_list) > 19:
        set_txbx_text(txbx_list[19], issued_year)
    if len(txbx_list) > 22:
        set_txbx_text(txbx_list[22], issued_day)
    if len(txbx_list) > 24:
        set_txbx_text(txbx_list[24], issued_month)
    if len(txbx_list) > 30:
        set_txbx_text(txbx_list[30], expiry_year)
    if len(txbx_list) > 33:
        set_txbx_text(txbx_list[33], expiry_day)
    if len(txbx_list) > 35:
        set_txbx_text(txbx_list[35], expiry_month)

    # ── Sertifikat 2 (o'ng/ikkinchi nusxa) ───────────────────────────────────
    # Word da ikki nusxa bor: birinchisi 0-20, ikkinchisi 21-40 oralig'i
    # Offset: birinchi nusxa indeksiga +21 qo'shsak ikkinchi nusxa indeksiga tushamiz
    SECOND_OFFSET = 21  # 9+21=30? — keling tekshirib chiqaylik

    # TextBox indekslari ikkinchi nusxa uchun (Word da ikkinchi nusxa ham mavjud):
    # TextBox[9+21=30] → ikkinchi Sertifikat ID joyida emas.
    # Aniqroq: birinchi nusxa 0-20, ikkinchi nusxa 21-40
    # TextBox[9]  → cert_id 1,  TextBox[9+21=30]  → yil (expiry)
    # Shunday ekan, ikkinchi nusxa indekslari boshqacha.
    # Keling, tahlil natijalariga ko'ra aniqlaylik:
    # TextBox[21..40] - ikkinchi nusxa:
    #   [21] '  '  → cert_id 2 (?), [22..29] sanalar
    #   Aslida TextBox[9] → 1-cert ID (bo'sh), TextBox[9+21=30] → '2027' (expiry yil 1)
    # => Ikkinchi nusxa ham xuddi shunday tuzilishda, faqat offset bilan

    # Korrekt ikkinchi nusxa mapping (tahlilga asosan):
    # TextBox[0..20]  = birinchi nusxa
    # TextBox[21..40] = ikkinchi nusxa
    # Birinchi nusxada:
    #   [9]  = cert_id
    #   [15] = talaba ismi
    #   [19] = boshlanish yili
    #   [22] = boshlanish kuni
    #   [24] = boshlanish oyi
    #   [30] = tugash yili
    #   [33] = tugash kuni
    #   [35] = tugash oyi
    # Ikkinchi nusxadagi offset: har birini +21 qilsak:
    #   [9+21=30]  = expiry_year (✗ bu allaqachon expiry yili!)
    # Bu to'g'ri kelmaydi. Demak offset boshqacha.
    # Tahlildan: TextBox[30]='2027', TextBox[31]=' yil' — bu expiry yili 1-nusxa uchun.
    # Ikkinchi nusxa uchun: 0-20 birinchi, 21-40 ikkinchi deb hisoblasak:
    #   [21+9=30] → '2027' — bu birinchi nusxa expiry yili!
    # Demak ikkinchi nusxa ALOHIDA hisoblanmagan, ikkisi bir xil tuzilish.
    # Haqiqiy mapping:
    #   Birinchi sertifikat: [9, 15, 19, 22, 24, 30, 33, 35]
    #   Ikkinchi sertifikat: [9+21=30? — yo'q, bu expiry yili]
    # Keling oddiy usul: XML ni to'g'ridan-to'g'ri tekshirib chiqamiz
    # TextBox[0..40] dan 21 ta birinchi nusxa, 20 ta ikkinchi nusxa deyish mumkin emas.
    # Hozircha faqat birinchi nusxani to'ldiramiz.

    # ── XML ni hujjatga qayta yuklash ─────────────────────────────────────────
    # python-docx da XML ni to'g'ridan manipulyatsiya qilish eng ishonchli usul
    # body elementini yangilaymiz
    body = doc.element.body
    # Body XML ni almashtiramiz
    new_body_xml = etree.tostring(root, encoding='unicode')

    # Hujjatni xotiraga saqlash
    docx_io = BytesIO()
    doc.save(docx_io)
    docx_io.seek(0)

    # Qayta ochamiz va body ni o'zgartiramiz (python-docx XML manipulyatsiyasi)
    doc2 = DocxDocument(docx_io)
    body2 = doc2.element.body

    # XML root dan body elementini olamiz va almashtirish qilmaymiz —
    # buning o'rniga to'g'ridan-to'g'ri run textlarini o'zgartiramiz

    # ── Eng ishonchli usul: to'g'ridan run textlarini topib o'zgartirish ─────
    # Hujjatni qayta ochib XML ustida ishlash
    doc3 = DocxDocument(template_path)
    root3 = doc3.element.body

    # Barcha w:t elementlarini olamiz
    all_t = root3.findall(f'.//{{{W_NS}}}t')

    # Template da bo'sh joy bor joylarni topib o'zgartirish o'rniga,
    # Python-docx ning to'g'ridan XMLni o'zgartirish usulini qo'llaymiz:
    txbx3 = root3.findall(f'.//{{{WPS_NS}}}txbx')

    def replace_txbx_runs(txbx, new_text):
        """Replace text in a textbox's first paragraph, keeping formatting."""
        paras = txbx.findall(f'.//{{{W_NS}}}p')
        if not paras:
            return
        para = paras[0]
        runs = para.findall(f'{{{W_NS}}}r')

        # Save first run's rPr for formatting
        first_rpr = None
        if runs:
            first_rpr = runs[0].find(f'{{{W_NS}}}rPr')

        # Remove all runs
        for r in runs:
            para.remove(r)

        # Create single new run
        new_run = etree.SubElement(para, f'{{{W_NS}}}r')
        if first_rpr is not None:
            new_run.insert(0, copy.deepcopy(first_rpr))
        t_el = etree.SubElement(new_run, f'{{{W_NS}}}t')
        t_el.text = new_text
        t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    if len(txbx3) > 9:
        replace_txbx_runs(txbx3[9], certificate.certificate_id)
    if len(txbx3) > 15:
        replace_txbx_runs(txbx3[15], student_name.upper())
    if len(txbx3) > 19:
        replace_txbx_runs(txbx3[19], issued_year)
    if len(txbx3) > 22:
        replace_txbx_runs(txbx3[22], issued_day)
    if len(txbx3) > 24:
        replace_txbx_runs(txbx3[24], issued_month)
    if len(txbx3) > 30:
        replace_txbx_runs(txbx3[30], expiry_year)
    if len(txbx3) > 33:
        replace_txbx_runs(txbx3[33], expiry_day)
    if len(txbx3) > 35:
        replace_txbx_runs(txbx3[35], expiry_month)

    # ── Hujjatni xotiraga saqlash ─────────────────────────────────────────────
    docx_output_io = BytesIO()
    doc3.save(docx_output_io)
    docx_output_io.seek(0)

    # ── Media ga saqlash ──────────────────────────────────────────────────────
    os.makedirs(os.path.join(settings.MEDIA_ROOT, 'certificates', 'words'), exist_ok=True)

    word_file_name = f"{certificate.certificate_id}.docx"
    word_rel_path = f"certificates/words/{word_file_name}"
    word_full_path = os.path.join(settings.MEDIA_ROOT, word_rel_path)

    with open(word_full_path, 'wb') as f:
        f.write(docx_output_io.read())

    return word_rel_path


# ─────────────────────────────────────────────────────────────────────────────
# 4. Main entry point — called from views.py
# ─────────────────────────────────────────────────────────────────────────────
def generate_pdf_and_qr(certificate):
    """
    Generate QR code, PDF certificate (overlay on template), and
    Word certificate (.docx) for the given certificate object.
    Sets certificate.qr_code_image and certificate.pdf_file (unsaved).
    Also creates media/certificates/words/<id>.docx on disk.
    """
    os.makedirs(os.path.join(settings.MEDIA_ROOT, 'certificates', 'pdfs'), exist_ok=True)
    os.makedirs(os.path.join(settings.MEDIA_ROOT, 'certificates', 'qrs'),  exist_ok=True)
    os.makedirs(os.path.join(settings.MEDIA_ROOT, 'certificates', 'words'), exist_ok=True)

    # 1. QR Code
    _generate_qr(certificate)

    # 2. PDF (needs QR saved to disk first)
    _generate_pdf(certificate)

    # 3. Word
    try:
        _generate_word(certificate)
    except Exception as e:
        # Word generatsiya muvaffaqiyatsiz bo'lsa, PDF ni bloklamaymiz
        import logging
        logging.getLogger(__name__).warning(f"Word generation failed for {certificate.certificate_id}: {e}")
